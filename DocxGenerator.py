from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date, datetime
import os

def generateDocx(result):
    folder_name = 'economic_analisis_2024'
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    day = date.today()
    full_day = datetime.today()

    title = doc.add_paragraph(f'\nAnálisis de la Economía Argentina dia {day}')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(24)

    doc.add_paragraph('\n\n')
    content = doc.add_paragraph(result)
    content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Guardar el documento
    doc_path = os.path.join(folder_name, f'Analisis_Economia_Argentina_{full_day}.docx')
    doc.save(doc_path)
